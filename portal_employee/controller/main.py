# -*- coding: utf-8 -*-

from odoo import http,_
from odoo.http import request
from odoo.addons.portal.controllers.portal import CustomerPortal
from odoo.exceptions import AccessError, MissingError,ValidationError,UserError
import base64
import logging

from datetime import datetime, date, timedelta, time
from psycopg2 import IntegrityError
from odoo.addons.auth_signup.models.res_partner import SignupError, now
from odoo.addons.website.controllers.main import Home
from docx import Document   #Import the Document
from docx.shared import Inches   #Import Inches, used to set width, height etc
from docx.shared import Pt   #Import Pt, used to font etc
import pytz

_logger = logging.getLogger(__name__)

class Home(Home):
    @http.route()
    def index(self, *args, **kw):
        values = {}
        employee_rec = self.get_employee()
        res = super(Home, self).index(*args, **kw)
        if not employee_rec:
            return res
        else:
            leave_data = request.env['hr.leave'].sudo().search([('employee_id', '=', employee_rec.id)])
            letter_request_data = request.env['employee.letter.request'].sudo().search([('employee_id', '=', employee_rec.id)])
            payslip_data = request.env['hr.payslip'].sudo().search([('employee_id', '=', employee_rec.id)])
            # if not current_login_user.has_group('hr.group_hr_manager'):
            #     is_manager = True
            # else:
            #     is_manager = False
            team_recs = request.env['hr.employee'].sudo().search([('parent_id', '=', employee_rec.id)])
            team_leaves_lst = []
            team_reqst_lst = []
            for team in team_recs:
                leave_recs = request.env['hr.leave'].sudo().search([('employee_id', '=', team.id)])
                team_leaves_lst.append({leave_recs})
                reqst_recs = request.env['employee.letter.request'].sudo().search([('employee_id', '=', team.id)])
                team_reqst_lst.append({reqst_recs})
            values.update({
                'employee': employee_rec,
                'leave_data': leave_data,
                'leave_count': len(leave_data),
                'payslip_data': payslip_data,
                'payslip_count': len(payslip_data),
                'letter_request_data': letter_request_data,
                'letter_request_count': len(letter_request_data),
                'team_data': team_recs,
                'team_count': len(team_recs),
                'team_leave_data': team_leaves_lst,
                'team_letter_data': team_reqst_lst,
            })
            return request.render("portal_employee.employee_portal_dashboard", values)
        return res

    @http.route('/editdetails', type='http', auth='user', website=True)
    def navigate_to_edit_employee_page(self):
        employee_rec = self.get_employee()
        return http.request.render('portal_employee.self_service_update_details', {
          'employee': employee_rec})

    @http.route(['/update_employee_form'], type='http', auth='user', website=True)
    def update_employee_form(self, **post):
        print("post..",post)
        qcontext = request.params.copy()
        if post.get('employee'):
             try:
                employee_rec = self.get_employee()
                response = request.render('portal_employee.self_service_update_details', qcontext)
                response.qcontext['employee'] = employee_rec
                response.headers['X-Frame-Options'] = 'DENY'
                employee_rec = request.env['hr.employee'].sudo().browse(int(post.get('employee')))
                employee_rec.address_home_id.update({
                    'street': post.get('address_street') if post.get('address_street') else employee_rec.address_home_id.street,
                    'street2': post.get('address_street2') if post.get('address_street2') else employee_rec.address_home_id.street2,
                    'city': post.get('address_city') if post.get('address_city') else employee_rec.address_home_id.city,
                    'zip': post.get('address_zip') if post.get('address_zip') else employee_rec.address_home_id.zip,
                    'state_id': int(post.get('state_id')) if post.get('state_id') else employee_rec.address_home_id.state_id.id,
                })
                employee_rec.bank_account_id.update({
                    'acc_number': post.get('acc_number') if post.get('acc_number') else employee_rec.bank_account_id.acc_number,
                    'bank_id': post.get('bank_id') if post.get('bank_id') else employee_rec.bank_account_id.bank_id.id,
                    'acc_holder_name': post.get('acc_holder_name') if post.get('acc_holder_name') else employee_rec.bank_account_id.acc_holder_name,
                })
                employee_rec.update({
                    'identification_id': post.get('identification_id') if post.get('identification_id') else employee_rec.identification_id,
                    'passport_id': post.get('passport_id') if post.get('passport_id') else employee_rec.passport_id,
                })
                attached_files = request.httprequest.files.getlist('attachments')
                for attachment in attached_files:
                    # attached_file = attachment.read()
                    request.env['ir.attachment'].sudo().create({
                                'name': attachment.filename,
                                'res_model': 'hr.employee',
                                'res_id': employee_rec.id,
                                'mimetype': 'binary',
                                # 'datas_fname': attachment.filename,
                                # 'datas': attached_file.encode('base64'),
                                'datas': base64.b64encode(attachment.read()),
                            })
                return request.render('portal_employee.request_submit_thankyou', {'update_profile':True})
                # return request.redirect('/request-submit-thank-you')
             except Exception as e:
                 value = {}
                 value['leave_data'] = request.env['ir.ui.view']._render_template("portal_employee.update_employee_form", {
                 })
                 qcontext['error'] = e.args[0]
                 request.env.cr.rollback()
                 return value
        else:
            response = request.render("website.homepage", qcontext)
        return response

    def get_employee(self):
        employee_rec = request.env['hr.employee'].sudo().search([('user_id', '=', request.env.user.id)], limit=1)
        return employee_rec

    # @http.route(['/website_payment/transaction/<string:reference>/<string:amount>/<string:currency_id>',
    #             '/website_payment/transaction/v2/<string:amount>/<string:currency_id>/<path:reference>',
    #             '/website_payment/transaction/v2/<string:amount>/<string:currency_id>/'
    #             '<path:reference>/<int:partner_id>'], type='json', auth='public')
    # def transaction(self, acquirer_id, reference, amount, currency_id, partner_id=False, **kwargs):
    #     acquirer = request.env['payment.acquirer'].browse(acquirer_id)
    #     order_id = kwargs.get('order_id')
    #

    @http.route(['/leaverequest',
                '/leaverequest/<string:emp_model>',
                ], type='http', auth='user', website=True)
    def navigate_to_leave_page(self, emp_model=False, **kwargs):
        is_emp = False
        if emp_model:
            employee_rec = emp_model
        else:
            employee_rec = self.get_employee()
            is_emp = True
        return http.request.render('portal_employee.self_service_leave_form', {
          'employee': employee_rec,'is_emp':is_emp})

    @http.route(['/submit_leave_request'], type='http', auth='user', csrf=False,website=True)
    def submit_leave_request(self, **post):
        qcontext = request.params.copy()
        if post.get('employee'):
            dates = post.get('dates')
            dates = dates.split('-')
            date_from = dates[0]
            date_from = "%s 08:00:00" %date_from
            date_from = datetime.strptime(date_from, '%m/%d/%Y %H:%M:%S')
            date_to = dates[1].lstrip()
            date_to = "%s 17:00:00" %date_to
            date_to = datetime.strptime(date_to, '%m/%d/%Y %H:%M:%S')
            local = pytz.timezone(request.env.user.tz or pytz.utc)
            date_from = datetime.strftime(pytz.utc.localize(date_from).astimezone(local),"%m/%d/%Y %H:%M:%S")
            date_from = datetime.strptime(date_from, '%m/%d/%Y %H:%M:%S')
            date_to = datetime.strftime(pytz.utc.localize(date_to).astimezone(local),"%m/%d/%Y %H:%M:%S")
            date_to = datetime.strptime(date_to, '%m/%d/%Y %H:%M:%S')
            try:
                employee_rec = self.get_employee()
                response = request.render('portal_employee.self_service_leave_form', qcontext)
                response.qcontext['employee'] = employee_rec
                response.headers['X-Frame-Options'] = 'DENY'
                leave_rec = request.env['hr.leave'].create({
                            'employee_id': int(post.get('employee')),
                            'request_date_from': date_from,
                            'date_from': date_from,
                            'holiday_type':'employee',
                            'request_date_to': date_to,
                            'date_to': date_to,
                            'name': post.get('description'),
                            'holiday_status_id': int(post.get('leave_type')),
                            # 'number_of_days': 0
                        })
                leave_rec._compute_date_from_to()
                leave_rec._compute_number_of_days()
                file = post.get('attachments')
                Attachments = request.env['ir.attachment']
                if file:
                    Attachments.create({
                        'name': file.filename,
                        'type': 'binary',
                        'datas': base64.b64encode(file.read()),
                        'res_model': "hr.leave",
                        'res_id': leave_rec.id
                    })
                return request.redirect('/request-submit-thank-you')
            except Exception as e:
                qcontext['error'] = e.args[0]
                request.env.cr.rollback()
                return response
        else:
            response = request.render("website.homepage", qcontext)
        return response

    @http.route(['/approve/leaverequest/<tl>'], type="http", auth="user", methods=['post'], website=True)
    def employee_approve_leave(self,  **kwargs):
        leave_id = kwargs.get('tl')
        qcontext = request.params.copy()
        if leave_id:
            try:
                Leave = request.env['hr.leave']
                leave_rec = Leave.sudo().browse(int(leave_id))
                leave_rec.action_approve()
                return request.redirect('/request-submit-thank-you')
            except Exception as e:
                qcontext['error'] = e.args[0]
                request.env.cr.rollback()
        employee_rec = self.get_employee()
        response = request.render("portal_employee.request_submit_feedback", qcontext)
        response.qcontext['employee'] = employee_rec
        response.headers['X-Frame-Options'] = 'DENY'
        return response

    @http.route(['/refuse/leaverequest/<tl>'], type="http", auth="user", methods=['post'], website=True)
    def employee_refuse_leave(self, **kwargs):
        leave_id = kwargs.get('tl')
        qcontext = request.params.copy()
        if leave_id:
            try:
                Leave = request.env['hr.leave']
                leave_rec = Leave.sudo().browse(int(leave_id))
                leave_rec.action_refuse()
                return request.redirect('/request-submit-thank-you')
            except Exception as e:
                qcontext['error'] = e.args[0]
                request.env.cr.rollback()
            # except UserError as e:
            #     qcontext['error'] = e.args[0]
            # except AccessError as ae:
            #     qcontext['error'] = ae.args[0]
            # except IntegrityError as inst:
            #     qcontext['error'] = inst.args[0]
        # response = request.render("website.homepage", qcontext)
        employee_rec = self.get_employee()
        response = request.render("portal_employee.request_submit_feedback", qcontext)
        response.qcontext['employee'] = employee_rec
        response.headers['X-Frame-Options'] = 'DENY'
        return response

    @http.route('/letterrequest', type='http', auth='user', website=True)
    def navigate_to_letter_page(self):
        employee_rec = self.get_employee()
        return http.request.render('portal_employee.self_service_letter_form', {
          'employee': employee_rec})

    @http.route(['/submit_letter_request'], type='http', auth='user', csrf=False,website=True)
    def submit_letter_request(self, **post):
        if post.get('employee'):
            qcontext = request.params.copy()
            try:
                ltype = post.get('ltype')
                position = post.get('position')
                if position == 'Iqama Position':
                    position = 'iqama_position'
                else:
                    position = 'company_position'
                description = post.get('description')
                salary = post.get('salary')
                chamber_of_commerce = post.get('chamberofcommerce')
                letter_req_rec = request.env['employee.letter.request'].sudo().create({
                        'employee_id': int(post.get('employee')),
                        'ltype': ltype,
                        'position': position,
                        'salary': salary,
                        'chamber_of_commerce': chamber_of_commerce,
                        'description': description,
                        'state': 'draft',
                    })
                if not letter_req_rec.chamber_of_commerce:
                    email_values = {'reply_to': request.env.company.email or ''}
                    if ltype=='arabic':
                        template = request.env.ref('portal_employee.email_template_for_employee_letter_request_arabic')
                    else:
                        template = request.env.ref('portal_employee.email_template_for_employee_letter_request')
                    template.sudo().send_mail(letter_req_rec.id, email_values=email_values, force_send=True, raise_exception=True)
                else:
                    Attachments = request.env['ir.attachment']
                    template = request.env.ref('portal_employee.email_template_for_regional_officers')
                    document = Document()
                    section = document.sections[0]   # Create a section
                    sec_header = section.header   # Create header
                    header_tp = sec_header.add_paragraph()  # Add a paragraph in the header, you can add any anything in the paragraph
                    header_run = header_tp.add_run()   # Add a run in the paragraph. In the run you can set the values
                    rml_header = "\t\t\t\t Letter Request"
                    header_run.add_text(rml_header)
                    header_run.font.size =  Pt(20)
                    header_run.add_text("\n\n\n")
                    header_run.font.size =  Pt(18)
                    table_main = document.add_table(rows=1, cols=2)
                    table_main.allow_autofit = True
                    tx_cells = table_main.rows[0].cells
                    tb_cell_run = tx_cells[0].add_paragraph().add_run()
                    emp_name = letter_req_rec.employee_id.name
                    tb_cell_run.add_text(emp_name)
                    tb_cell_run.font.size =  Pt(18)
                    tx_cells[0].width = Inches(6)
                    date = letter_req_rec.create_date.strftime('%Y-%m-%d')
                    date = "\t\t\t"+str(date)
                    tb_cell_run = tx_cells[1].add_paragraph().add_run()
                    tb_cell_run.add_text(date)
                    tb_cell_run.font.size =  Pt(16)
                    tx_cells[1].width = Inches(6)
                    details = document.add_paragraph()
                    details_run = details.add_run()
                    details_run.text = "\n\n\t\t Messrs. / Heath of Saudi Engineers\n\n"
                    details_run.font.size = Pt(18)
                    details = document.add_paragraph()
                    details_run = details.add_run()
                    details_run.text = "Greetings,\n"
                    details_run.font.size = Pt(20)
                    details = document.add_paragraph()
                    details_run = details.add_run()
                    text = "We, the Moment of Sunrise Trading Corporation, inform you that "+ str(emp_name) + ", of Jordanian nationality, with border number No. " + str(letter_req_rec.name) +" , has been working for us since" + str(date) +" with the position of (computer programmer) and is still on the job. This certificate was given at his request to register with the authority and obtain residency without any responsibility on the part of the institution.\n\n\n"
                    details_run.text = text
                    details_run.font.size = Pt(18)
                    details = document.add_paragraph()
                    details_run = details.add_run()
                    text = "Accept my sincere greetings and appreciation,"
                    details_run.text = text
                    details_run.font.size = Pt(16)
                    details = document.add_paragraph()
                    details_run = details.add_run()
                    text = "General Director"
                    details_run.text = text
                    details_run.font.size = Pt(20)
                    document.save("/tmp/letter_request.docx")  # Save the file
                    file_doc1_base64 = base64.b64encode(open('/tmp/letter_request.docx', 'rb').read())
                    attachment = Attachments.create({'name': 'Letter Request for ' + letter_req_rec.employee_id.name,
                                               'datas': file_doc1_base64,
                                               'res_id': letter_req_rec.id,
                                               # 'type': 'binary',
                                               'mimetype': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                                                'res_model': 'employee.letter.request',
                                                   })
                    email_list = [line.name for line in letter_req_rec.employee_id.region_id.region_line_ids]
                    email_to = ','.join(email_list)
                    email_values = {'email_to': email_to,
                                    'subject': 'Letter Request for-' + str(letter_req_rec.employee_id.name),
                                    'attachment_ids': [(6, 0, attachment.ids)],
                                }
                    template.sudo().send_mail(letter_req_rec.id, email_values=email_values, force_send=True, raise_exception=True)
                return request.redirect('/request-submit-thank-you')
            except Exception as e:
                response = request.render('portal_employee.self_service_letter_form', qcontext)
                response.qcontext['employee'] = self.get_employee()
                response.headers['X-Frame-Options'] = 'DENY'
                qcontext['error'] = e.args[0]
                request.env.cr.rollback()
                return response
        return request.redirect('/')

    @http.route(['/approve/letterrequest/<tr>'], type="http", auth="user", methods=['post'], website=True)
    def employee_approve_letter(self,  **kwargs):
        letter_id = kwargs.get('tr')
        qcontext = request.params.copy()
        if letter_id:
            try:
                Letter = request.env['employee.letter.request']
                letter_rec = Letter.sudo().browse(int(letter_id))
                letter_rec.action_approve()
                return request.redirect('/request-submit-thank-you')
            except Exception as e:
                qcontext['error'] = e.args[0]
                request.env.cr.rollback()
        # response = request.render("website.homepage", qcontext)
        employee_rec = self.get_employee()
        response = request.render("portal_employee.request_submit_feedback", qcontext)
        response.qcontext['employee'] = employee_rec
        response.headers['X-Frame-Options'] = 'DENY'
        return response

    @http.route(['/refuse/letterrequest/<tr>'], type="http", auth="user", methods=['post'], website=True)
    def employee_refuse_letter(self, **kwargs):
        letter_id = kwargs.get('tr')
        qcontext = request.params.copy()
        if letter_id:
            try:
                Letter = request.env['employee.letter.request']
                letter_rec = Letter.sudo().browse(int(letter_id))
                letter_rec.action_refuse()
                return request.redirect('/request-submit-thank-you')
            except Exception as e:
                qcontext['error'] = e.args[0]
                request.env.cr.rollback()
        # response = request.render("website.homepage", qcontext)
        employee_rec = self.get_employee()
        response = request.render("portal_employee.request_submit_feedback", qcontext)
        response.qcontext['employee'] = employee_rec
        response.headers['X-Frame-Options'] = 'DENY'
        return response


    # @http.route(['/refuse/leaverequest/<tl>'], type="http", auth="user", methods=['post'], website=True)
    # def submit_rating(self, leave_id=False, **kwargs):
    #     leave_id = kwargs.get('tl')
    #     qcontext = request.params.copy()
    #     if leave_id:
    #         try:
    #             Leave = request.env['hr.leave']
    #             leave_rec = Leave.sudo().browse(int(leave_id))
    #             leave_rec.action_refuse()
    #             return request.redirect('/request-submit-thank-you')
    #         except UserError as e:
    #             qcontext['error'] = e.args[0]
    #         except IntegrityError as inst:
    #             qcontext['error'] = inst.args[0]
    #     employee_rec = self.get_employee()
    #     response = request.render("portal_employee.request_submit_feedback", qcontext)
    #     response.qcontext['employee'] = employee_rec
    #     response.headers['X-Frame-Options'] = 'DENY'
    #     return response

# class CustomerPortal(CustomerPortal):
    # @http.route(['/leaverequest'], type='http', auth="user", website=True)
    # def gallery_categories(self, employee=None, **kw):
    #     values = {}
    #     values.update({
    #         'employee': employee,
    #     })
    #     return request.render("portal_employee.leave_request_template_id", values)

    # @http.route(['/employee/profile/update_json'], type='json', auth="public", methods=['POST'], website=True)
    # def employee_profile_update(self, **kw):
    #     value = {}
    #     gender = False
    #     marital_status = False
    #     if kw.get('gender') == 'Male':
    #         gender = 'male'
    #     if kw.get('gender') == 'Female':
    #         gender = 'female'
    #     if kw.get('gender') == 'Other':
    #         gender = 'other'
    #     if kw.get('marital_status') == 'Single':
    #         marital_status = 'single'
    #     if kw.get('marital_status') == 'Married':
    #         marital_status = 'married'
    #     current_login_user = request.env.user
    #     employee_id = request.env['hr.employee'].sudo().search([('user_id', '=', current_login_user.id)], limit=1)
    #     country_id = request.env['res.country'].sudo().search([('id', '=', int(kw.get('private_country_id')))], limit=1)
    #     state_id = request.env['res.country.state'].sudo().search([('id', '=', int(kw.get('private_state_id')))], limit=1)
    #     if kw.get('image'):
    #         image = BytesIO(b64decode(kw.get('image').split(',')[1]))
    #         employee_id.update({'image_1920': base64.b64encode(image.read()).decode('utf-8')})
    #     employee_id.update({
    #             'name': kw.get('name'),
    #             'identification_id': kw.get('identification_id'),
    #             'passport_id': kw.get('passport_id'),
    #             'birthday': kw.get('birthday'),
    #             'gender': gender,
    #             'marital': marital_status,
    #         })
    #     employee_id.address_home_id.update({
    #             'street': kw.get('private_street'),
    #             'street2': kw.get('private_street2'),
    #             'zip': kw.get('private_zip'),
    #             'state_id': state_id.id,
    #             'country_id': country_id.id,
    #         })
    #     value['render_employee_data'] = request.env['ir.ui.view']._render_template("portal_employee.portal_employee_data", {
    #         })
    #     value['render_employee_form'] = request.env['ir.ui.view']._render_template("portal_employee.portal_employee_form", {
    #         })
    #     return value
    #
    # @http.route(['/manager/employee/profile/update_json'], type='json', auth="public", methods=['POST'], website=True)
    # def manager_employee_profile_update(self, **kw):
    #     value = {}
    #     employee_id = request.env['hr.employee'].sudo().browse(int(kw.get('employee_id')))
    #     value['render_employee_form'] = request.env['ir.ui.view']._render_template("portal_employee.manager_portal_employee_form", {'employee_id': employee_id})
    #     return value
    #
    # @http.route(['/save/employee/profile/update_json'], type='json', auth="public", methods=['POST'], website=True)
    # def save_employee_profile_update(self, **kw):
    #     value = {}
    #     gender = False
    #     marital_status = False
    #     if kw.get('gender') == 'Male':
    #         gender = 'male'
    #     if kw.get('gender') == 'Female':
    #         gender = 'female'
    #     if kw.get('gender') == 'Other':
    #         gender = 'other'
    #     if kw.get('marital_status') == 'Single':
    #         marital_status = 'single'
    #     if kw.get('marital_status') == 'Married':
    #         marital_status = 'married'
    #     employee_id = request.env['hr.employee'].sudo().search([('id', '=', int(kw.get('emp_id')))], limit=1)
    #     country_id = request.env['res.country'].sudo().search([('id', '=', int(kw.get('private_country_id')))], limit=1)
    #     state_id = request.env['res.country.state'].sudo().search([('id', '=', int(kw.get('private_state_id')))], limit=1)
    #     if kw.get('image'):
    #         image = BytesIO(b64decode(kw.get('image').split(',')[1]))
    #         employee_id.update({'image_1920': base64.b64encode(image.read()).decode('utf-8')})
    #     employee_id.update({
    #             'name': kw.get('name'),
    #             'identification_id': kw.get('identification_id'),
    #             'passport_id': kw.get('passport_id'),
    #             'birthday': kw.get('birthday'),
    #             'gender': gender,
    #             'marital': marital_status,
    #         })
    #     employee_id.address_home_id.update({
    #             'street': kw.get('private_street'),
    #             'street2': kw.get('private_street2'),
    #             'zip': kw.get('private_zip'),
    #             'state_id': state_id.id,
    #             'country_id': country_id.id,
    #         })
    #     value['render_o_organizational_chart'] = request.env['ir.ui.view']._render_template("portal_employee.o_organizational_chart", {})
    #     return value
    #
    # @http.route(['/leave/update_json'], type='json', auth="public", methods=['POST'], website=True)
    # def leave_update_json(self, **kw):
    #     value = {}
    #     current_login_user = request.env.user
    #     employee_id = request.env['hr.employee'].sudo().search([('user_id', '=', current_login_user.id)], limit=1)
    #     holiday = request.env['hr.leave'].sudo()
    #     holiday.create({
    #             'employee_id': employee_id.id,
    #             'request_date_from': kw.get('from_date'),
    #             'request_date_to': kw.get('to_date'),
    #             'name': kw.get('description'),
    #             'holiday_status_id': int(kw.get('holiday_status_id')),
    #             'number_of_days': int(kw.get('number_of_days'))
    #         })
    #     value['leave_data'] = request.env['ir.ui.view']._render_template("portal_employee.leave_data", {
    #         })
    #     value['render_leave_form'] = request.env['ir.ui.view']._render_template("portal_employee.leave_form", {
    #         })
    #     return value
    #
    # @http.route(['/calculate/number_of_days'], type='json', auth="public", methods=['POST'], website=True)
    # def calculate_number_of_days(self, **kw):
    #     current_login_user = request.env.user
    #     employee_id = request.env['hr.employee'].sudo().search([('user_id', '=', current_login_user.id)], limit=1)
    #     holiday = request.env['hr.leave'].sudo().search([('employee_id', '=', employee_id.id)], limit=1)
    #     if kw.get('from_date') and kw.get('to_date'):
    #         from_date = datetime.strptime(kw.get('from_date'), '%Y-%m-%d')
    #         to_date = datetime.strptime(kw.get('to_date'), '%Y-%m-%d')
    #         days = holiday._get_number_of_days(from_date, to_date, holiday.employee_id.id)['days']
    #         return days+1
    #
    # @http.route(['/my/payslip/report/<int:payslip_id>'], type='http', auth="public", website=True)
    # def print_payslip(self, payslip_id, access_token=None, report_type=None, download=False, **kw):
    #     current_login_user = request.env.user
    #     employee_id = request.env['hr.employee'].sudo().search([('user_id', '=', current_login_user.id)], limit=1)
    #     payslip = request.env['hr.payslip'].sudo().browse(int(payslip_id))
    #     if employee_id.id == payslip.employee_id.id:
    #         try:
    #             payslip = self._document_check_access('hr.payslip', payslip_id)
    #         except (AccessError, MissingError):
    #             return request.redirect('/my')
    #         if report_type in ('html', 'pdf', 'text'):
    #             return self._show_report(model=payslip, report_type=report_type, report_ref='hr_payroll_community.action_report_payslip', download=download)
    #
    # @http.route(['/approve/leave_update_json'], type='json', auth="public", methods=['POST'], website=True)
    # def approve_leave(self, **kw):
    #     value = {}
    #     leave_id = request.env['hr.leave'].sudo().browse(int(kw.get('leave_id')))
    #     leave_id.action_approve()
    #     value['leave_data'] = request.env['ir.ui.view']._render_template("portal_employee.leave_data", {
    #         })
    #     return value
    #
    # @http.route(['/refuse/leave_update_json'], type='json', auth="public", methods=['POST'], website=True)
    # def refuse_leave(self, **kw):
    #     value = {}
    #     leave_id = request.env['hr.leave'].sudo().browse(int(kw.get('leave_id')))
    #     leave_id.action_refuse()
    #     value['leave_data'] = request.env['ir.ui.view']._render_template("portal_employee.leave_data", {
    #         })
    #     return value

